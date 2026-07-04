from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "Dynalite_PDEG_Project_Summary.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.6, 4.9])
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Details"
    for cell in hdr:
        set_cell_shading(cell, TABLE_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    set_repeat_table_header(table.rows[0])
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = val
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.color.rgb = DARK_BLUE
    title.font.bold = True
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def build() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Dynalite PDEG Home Assistant Integration")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Project summary for the repository at ")
    subtitle.add_run(str(ROOT)).italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This repository contains a HACS-ready Home Assistant custom integration "
        "for Philips Dynalite lighting and automation systems connected through a "
        "PDEG Ethernet Gateway. It talks to the gateway over TCP, parses DyNet1 "
        "frames, keeps a live coordinator state model, and dynamically exposes "
        "Home Assistant entities for lighting, relays, blinds, HVAC, sensors, "
        "preset controls, and physical device diagnostics."
    )

    doc.add_heading("Repository Layout", level=1)
    add_key_value_table(
        doc,
        [
            ("Root", str(ROOT)),
            ("Integration domain", "dynalite_pdeg"),
            ("Main component", "custom_components/dynalite_pdeg"),
            ("HACS metadata", "hacs.json"),
            ("HA manifest", "custom_components/dynalite_pdeg/manifest.json"),
            ("Documentation", "README.md and docs/ reference files"),
        ],
    )

    doc.add_heading("Primary Capabilities", level=1)
    for item in [
        "Local TCP connection to a Philips Dynalite PDEG gateway.",
        "Home Assistant config flow, reconfigure flow, and dynamic platform setup.",
        "Automatic and manual discovery of Dynalite areas, channels, and physical devices.",
        "Custom Home Assistant sidebar panel for logical and physical configuration.",
        "System Builder XML import for logical areas and physical devices.",
        "Entity creation for lights, switches, covers, climate, binary sensors, sensors, selects, and buttons.",
        "Physical device sign-on polling with online/offline state tracking.",
        "Backup and restore through custom websocket commands.",
        "Developer/debug services including scan, select preset, set level, and raw DyNet frame send.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Core Architecture", level=1)
    doc.add_heading("TCP Client", level=2)
    doc.add_paragraph(
        "dynalite_client.py manages the persistent TCP connection to the PDEG. "
        "It reads fixed logical/physical frames and variable sign-on frames, "
        "validates checksums, reconnects on failure, and provides command helpers "
        "for presets, levels, occupancy, HVAC setpoints, sign-on polling, motion "
        "status requests, and other DyNet1 operations."
    )
    doc.add_heading("Coordinator", level=2)
    doc.add_paragraph(
        "coordinator.py is the integration's state engine. It owns AreaState, "
        "ChannelState, and PhysicalDevice data, decodes incoming frames, updates "
        "runtime state, persists structural changes through storage, and dispatches "
        "Home Assistant update signals to platform entities."
    )
    doc.add_heading("Home Assistant Platforms", level=2)
    add_key_value_table(
        doc,
        [
            ("light.py", "Dimmable and on/off channel lights."),
            ("switch.py", "Occupancy enable/disable switches and channel switches."),
            ("cover.py", "Preset-based blind covers and paired relay covers."),
            ("climate.py", "HVAC area entities with mode, fan, and setpoint support."),
            ("binary_sensor.py", "PIR motion, physical device connectivity, and device motion sensors."),
            ("sensor.py", "Temperature and illuminance sensors."),
            ("select.py", "Preset selector per lighting area."),
            ("button.py", "Save-preset button per lighting area."),
        ],
    )

    doc.add_heading("Configuration and UI", level=1)
    doc.add_paragraph(
        "Initial setup is handled by config_flow.py, which collects host, TCP port, "
        "and display name, then tests connectivity before creating the config entry. "
        "On load, __init__.py registers a gateway device, forwards platform setup, "
        "loads storage, starts the coordinator, registers services and websocket "
        "commands, and adds the custom sidebar panel."
    )
    doc.add_paragraph(
        "The custom panel in panel/dynalite-config-panel.js uses websocket.py. "
        "It exposes logical configuration for areas/channels/presets/PIR/HVAC/blinds "
        "and physical configuration for discovered devices, custom names, sign-on "
        "polling, XML import, lux/motion toggles, backup, and restore."
    )

    doc.add_heading("Persistence and Import", level=1)
    doc.add_paragraph(
        "storage.py uses Home Assistant's Store API to persist structural data per "
        "config entry. It saves areas, channel definitions, area types, names, "
        "preset counts, fade settings, PIR and temperature flags, HVAC mappings, "
        "curtains, physical devices, lux/motion flags, and sign-on interval. "
        "Live runtime state such as brightness, motion, temperature, and online "
        "status is intentionally rebuilt from bus traffic after startup."
    )
    doc.add_paragraph(
        "xml_parser.py supports Dynalite System Builder LogicalExport and DeviceExport "
        "files. Logical imports create areas, presets, and channels; device imports "
        "create physical module records keyed by device code and box number."
    )

    doc.add_heading("Services and Debugging", level=1)
    add_key_value_table(
        doc,
        [
            ("dynalite_pdeg.scan", "Poll an area/channel range for discovery."),
            ("dynalite_pdeg.select_preset", "Activate a preset in a Dynalite area."),
            ("dynalite_pdeg.set_level", "Set a specific channel level from 0 to 100 percent."),
            ("dynalite_pdeg.send_raw_frame", "Send raw DyNet hex frames, optionally recalculating checksum."),
        ],
    )

    doc.add_heading("Packaging Metadata", level=1)
    add_key_value_table(
        doc,
        [
            ("Manifest domain", "dynalite_pdeg"),
            ("Manifest name", "Philips Dynalite (PDEG)"),
            ("Version", "0.1.1"),
            ("Config flow", "Enabled"),
            ("IoT class", "local_push"),
            ("HACS minimum HA", "2024.1.0"),
            ("Python requirements", "None declared"),
        ],
    )

    doc.add_heading("Notable Observations", level=1)
    for item in [
        "The project has a strong separation between protocol I/O, coordinator state, and Home Assistant entity platforms.",
        "The integration is significantly more capable than a basic light adapter; it models both logical areas and physical hardware.",
        "There are no tests or CI files visible in the repository.",
        "README.md references an MIT license file, but no LICENSE file was visible in the root.",
        "README.md discusses PDEG ports around 50000/50001, while const.py currently sets DEFAULT_PORT to 12345.",
        "Some terminal output showed encoding artifacts in comments and documentation strings; the files may need UTF-8 display or cleanup.",
        "The raw frame service directly accesses the client's private writer, which works but is a maintenance smell.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Overall Assessment", level=1)
    doc.add_paragraph(
        "This is a mature custom Home Assistant integration with a practical architecture "
        "for real Dynalite installations: persistent TCP communication, central state "
        "coordination, dynamic entity creation, a custom admin panel, and import/backup "
        "tools. The main opportunities are project hygiene and hardening: adding tests, "
        "CI, a license file, documentation consistency checks, and small cleanup around "
        "encoding and private attribute access."
    )

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
